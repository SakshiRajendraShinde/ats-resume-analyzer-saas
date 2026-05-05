from docx import Document
import os

os.makedirs("resumes", exist_ok=True)

data = [
    ("resume_rahul.docx", "Python SQL Machine Learning Pandas NumPy Pune Computer 2024"),
    ("resume_priya.docx", "Python Excel Data Analysis Mumbai IT 2023"),
    ("resume_amit.docx", "Mechanical AutoCAD SolidWorks Delhi 2022")
]

for filename, content in data:
    doc = Document()
    doc.add_heading("Resume", 0)
    doc.add_paragraph(content)
    doc.save(f"resumes/{filename}")

print("✅ Sample resumes created")